"""
Jindal Steel OHC - Unified App
(OPD Registration + Medicine Expiry Tracker + MedBot Chatbot)

Run: python jindal_ohc_app.py
Requirements: pip install kivy python-docx
"""

import os
import sys
import sqlite3
import subprocess
from datetime import datetime

from kivy.app import App
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.uix.scrollview import ScrollView
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.gridlayout import GridLayout
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.spinner import Spinner
from kivy.uix.button import Button
from kivy.uix.widget import Widget
from kivy.uix.popup import Popup
from kivy.graphics import Color, Rectangle, RoundedRectangle
from kivy.core.window import Window
from kivy.utils import get_color_from_hex
from kivy.clock import Clock
from kivy.metrics import dp

# ══════════════════════════════════════════════════════════════════════════
#  GLOBAL WINDOW SETTINGS (ek hi jagah set kiya, saari screens isko share
#  karti hain)
# ══════════════════════════════════════════════════════════════════════════
Window.size = (480, 820)
Window.clearcolor = (0.85, 0.93, 1, 1)   # Light blue background

# ── Shared Colours ───────────────────────────────────────────────────────
BLUE   = (0.05, 0.27, 0.55, 1)
RED    = (0.72, 0.11, 0.11, 1)
WHITE  = (1, 1, 1, 1)
DGRAY  = (0.3, 0.3, 0.3, 1)
GREEN  = (0.10, 0.55, 0.25, 1)
ORANGE = (0.85, 0.45, 0.05, 1)
BLACK  = (0, 0, 0, 1)


# ══════════════════════════════════════════════════════════════════════════
#  COMMON NAVIGATION BAR (teeno screens ke beech switch karne ke liye)
# ══════════════════════════════════════════════════════════════════════════
def build_nav_bar(screen):
    """screen = Screen instance jiska 'manager' baad me set hoga"""
    bar = BoxLayout(size_hint_y=None, height=dp(46), spacing=dp(4), padding=[dp(4), dp(4)])
    buttons = [
        ("OPD Form", "opd", BLUE),
        ("Medicines", "medicine", GREEN),
        ("MedBot", "medbot", (0.12, 0.23, 0.54, 1)),
    ]
    for label, name, color in buttons:
        btn = Button(text=label, bold=True, font_size=dp(12),
                     background_normal='', background_color=color, color=WHITE)
        btn.bind(on_release=lambda inst, n=name: setattr(screen.manager, 'current', n))
        bar.add_widget(btn)
    return bar


# ══════════════════════════════════════════════════════════════════════════
#  SCREEN 1 : OPD REGISTRATION FORM
# ══════════════════════════════════════════════════════════════════════════
def field(label_text, hint="", multiline=False, height=dp(38)):
    row = BoxLayout(orientation="horizontal", size_hint_y=None,
                     height=height if not multiline else dp(70), spacing=dp(6))
    lbl = Label(text=label_text, size_hint_x=0.38, color=DGRAY,
                font_size=dp(13), halign="right", valign="middle", bold=True)
    lbl.bind(size=lbl.setter("text_size"))
    inp = TextInput(hint_text=hint, size_hint_x=0.62, multiline=multiline,
                     font_size=dp(13), padding=[dp(8), dp(8)],
                     background_color=WHITE, foreground_color=BLACK,
                     cursor_color=BLUE,
                     size_hint_y=None, height=height if not multiline else dp(70))
    row.add_widget(lbl)
    row.add_widget(inp)
    return row, inp


def vital_field(label_text, unit_text, hint=""):
    row = BoxLayout(orientation="horizontal", size_hint_y=None, height=dp(42), spacing=dp(6))
    lbl = Label(text=label_text, size_hint_x=0.28, color=DGRAY,
                font_size=dp(13), bold=True, halign="right", valign="middle")
    lbl.bind(size=lbl.setter("text_size"))
    inp = TextInput(hint_text=hint, size_hint_x=0.45, font_size=dp(13),
                     padding=[dp(8), dp(8)], background_color=WHITE,
                     foreground_color=BLACK, size_hint_y=None, height=dp(38))
    unit = Label(text=unit_text, size_hint_x=0.27, color=RED,
                 font_size=dp(12), halign="left", valign="middle")
    unit.bind(size=unit.setter("text_size"))
    row.add_widget(lbl)
    row.add_widget(inp)
    row.add_widget(unit)
    return row, inp


class OPDScreen(Screen):
    def __init__(self, **kw):
        super().__init__(**kw)
        self.inputs = {}
        self._build()

    def _build(self):
        root = BoxLayout(orientation="vertical")
        root.add_widget(build_nav_bar(self))

        header = BoxLayout(orientation="vertical", size_hint_y=None, height=dp(90), padding=dp(8))
        with header.canvas.before:
            Color(*BLUE)
            self._header_rect = Rectangle(pos=header.pos, size=header.size)
        header.bind(pos=self._upd_hr, size=self._upd_hr)

        header.add_widget(Label(text="JINDAL STEEL — OHC", color=WHITE, font_size=dp(15),
                                 bold=True, size_hint_y=None, height=dp(26)))
        header.add_widget(Label(text="OCCUPATIONAL HEALTH CENTRE", color=(0.9, 0.9, 0.9, 1),
                                 font_size=dp(12), size_hint_y=None, height=dp(20)))
        header.add_widget(Label(text="OPD REGISTRATION CARD — Patratu (JH)",
                                 color=(0.75, 0.85, 1, 1), font_size=dp(11),
                                 size_hint_y=None, height=dp(18)))
        root.add_widget(header)

        scroll = ScrollView(size_hint=(1, 1))
        form = BoxLayout(orientation="vertical", size_hint_y=None, spacing=dp(6),
                          padding=[dp(12), dp(10), dp(12), dp(10)])
        form.bind(minimum_height=form.setter("height"))

        def section(title):
            lbl = Label(text=f"  {title}", size_hint_y=None, height=dp(28), color=WHITE,
                        font_size=dp(12), bold=True, halign="left", valign="middle")
            lbl.bind(size=lbl.setter("text_size"))
            with lbl.canvas.before:
                Color(*RED)
                rect = Rectangle()
                lbl.bind(pos=lambda w, v: setattr(rect, "pos", v),
                         size=lambda w, v: setattr(rect, "size", v))
            return lbl

        form.add_widget(section("ADMINISTRATIVE DETAILS"))
        for lname, hint, key in [
            ("Regd. No.", "e.g. OHC-2025-001", "regd_no"),
            ("Medical Officer", "Dr. Name", "doctor"),
            ("Date & Time", datetime.now().strftime("%d/%m/%Y  %H:%M"), "datetime"),
        ]:
            r, inp = field(lname, hint)
            if key == "datetime":
                inp.text = datetime.now().strftime("%d/%m/%Y  %H:%M")
            self.inputs[key] = inp
            form.add_widget(r)

        form.add_widget(section("PATIENT INFORMATION"))
        for lname, hint, key in [
            ("Name of Patient", "Full name", "name"),
            ("E Code / GP No.", "Employee / Guest ID", "ecode"),
            ("Age / Sex", "e.g.  32 / M", "age_sex"),
            ("Mob. No.", "10-digit mobile", "mobile"),
            ("Address / Deptt.", "Department / Address", "address"),
            ("Contractor", "Contractor name", "contractor"),
            ("Service", "e.g.  Employee / GC", "service"),
        ]:
            r, inp = field(lname, hint)
            self.inputs[key] = inp
            form.add_widget(r)

        form.add_widget(section("VITAL SIGNS"))
        for lname, unit, key in [
            ("Temp.", "°F", "temp"),
            ("SPO2", "%", "spo2"),
            ("PR", "/min", "pr"),
            ("BP", "mmHg", "bp"),
            ("Wt.", "Kg", "weight"),
        ]:
            r, inp = vital_field(lname, unit)
            self.inputs[key] = inp
            form.add_widget(r)

        form.add_widget(section("INVESTIGATION / NOTES"))
        r, inp = field("Investigation", "Lab tests, X-Ray…", multiline=True)
        self.inputs["investigation"] = inp
        form.add_widget(r)

        r, inp = field("Remarks", "Doctor's remarks…", multiline=True)
        self.inputs["remarks"] = inp
        form.add_widget(r)

        scroll.add_widget(form)
        root.add_widget(scroll)

        btn_bar = BoxLayout(size_hint_y=None, height=dp(54), spacing=dp(10), padding=[dp(10), dp(8)])
        btn_word = Button(text="Save as Word", background_color=BLUE, color=WHITE,
                           font_size=dp(13), bold=True)
        btn_print = Button(text="Print / Preview", background_color=GREEN, color=WHITE,
                            font_size=dp(13), bold=True)
        btn_clear = Button(text="Clear", background_color=ORANGE, color=WHITE,
                            font_size=dp(13), bold=True, size_hint_x=0.45)
        btn_word.bind(on_press=self.save_word)
        btn_print.bind(on_press=self.print_form)
        btn_clear.bind(on_press=self.clear_form)
        btn_bar.add_widget(btn_word)
        btn_bar.add_widget(btn_print)
        btn_bar.add_widget(btn_clear)
        root.add_widget(btn_bar)

        self.add_widget(root)

    def _upd_hr(self, inst, val):
        self._header_rect.pos = inst.pos
        self._header_rect.size = inst.size

    def _data(self):
        return {k: v.text.strip() for k, v in self.inputs.items()}

    def clear_form(self, *_):
        for k, v in self.inputs.items():
            v.text = ""
        self.inputs["datetime"].text = datetime.now().strftime("%d/%m/%Y  %H:%M")
        self._popup("Form Cleared", "All fields have been reset.")

    def save_word(self, *_):
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            self._popup("Missing Library",
                         "python-docx not installed.\n\nRun:\n  pip install python-docx", error=True)
            return

        d = self._data()
        doc = Document()
        for sec in doc.sections:
            sec.top_margin = Cm(1.5)
            sec.bottom_margin = Cm(1.5)
            sec.left_margin = Cm(2)
            sec.right_margin = Cm(2)

        def heading(text, level=1):
            p = doc.add_heading(text, level=level)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(13, 71, 141)
            return p

        def kv(label, value, bold_val=False):
            p = doc.add_paragraph()
            r1 = p.add_run(f"{label}: ")
            r1.bold = True
            r1.font.size = Pt(11)
            r2 = p.add_run(value or "—")
            r2.bold = bold_val
            r2.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(2)

        heading("JINDAL STEEL POWER LIMITED", level=1)
        heading("Occupational Health Centre — Patratu (JH)", level=2)
        heading("OPD REGISTRATION CARD", level=2)
        doc.add_paragraph()

        doc.add_heading("Administrative Details", level=3)
        kv("Regd. No.", d.get("regd_no"))
        kv("Medical Officer", d.get("doctor"))
        kv("Date & Time", d.get("datetime"))

        doc.add_paragraph()
        doc.add_heading("Patient Information", level=3)
        kv("Name of Patient", d.get("name"), bold_val=True)
        kv("E Code / GP No.", d.get("ecode"))
        kv("Age / Sex", d.get("age_sex"))
        kv("Mobile No.", d.get("mobile"))
        kv("Address / Deptt.", d.get("address"))
        kv("Contractor", d.get("contractor"))
        kv("Service", d.get("service"))

        doc.add_paragraph()
        doc.add_heading("Vital Signs", level=3)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Parameter"
        hdr[1].text = "Value"
        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

        vitals = [
            ("Temperature", f"{d.get('temp', '')}  °F"),
            ("SPO2", f"{d.get('spo2', '')}  %"),
            ("Pulse Rate", f"{d.get('pr', '')}  /min"),
            ("Blood Pressure", f"{d.get('bp', '')}  mmHg"),
            ("Weight", f"{d.get('weight', '')}  Kg"),
        ]
        for name, val in vitals:
            row = tbl.add_row().cells
            row[0].text = name
            row[1].text = val

        doc.add_paragraph()
        doc.add_heading("Investigation / Notes", level=3)
        doc.add_paragraph(d.get("investigation") or "—")
        doc.add_heading("Remarks", level=3)
        doc.add_paragraph(d.get("remarks") or "—")

        section_obj = doc.sections[0]
        section_obj.footer_distance = Cm(1)
        footer = section_obj.footer
        footer.is_linked_to_previous = False
        for para in list(footer.paragraphs):
            elem = para._element
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
        sig_para = footer.add_paragraph()
        sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_run = sig_para.add_run("Signature of Doctor: ____________________")
        sig_run.bold = True
        sig_run.font.size = Pt(11)
        sig_run.font.color.rgb = RGBColor(20, 20, 20)

        patient_name = (d.get("name") or "Patient").strip().replace(" ", "_") or "Form"
        fname = f"OPD_{patient_name}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        if os.path.isdir(desktop):
            save_path = os.path.join(desktop, fname)
            location_label = "Desktop"
        else:
            save_path = os.path.join(os.getcwd(), fname)
            location_label = os.getcwd()

        try:
            doc.save(save_path)
            if os.path.isfile(save_path):
                size_kb = round(os.path.getsize(save_path) / 1024, 1)
                self._try_open(save_path)
                self._popup_save_success(fname, save_path, location_label, size_kb)
            else:
                self._popup("Save Failed", "File was not created on disk.\nCheck folder permissions.", error=True)
        except Exception:
            fallback_path = os.path.join(os.getcwd(), fname)
            try:
                doc.save(fallback_path)
                if os.path.isfile(fallback_path):
                    size_kb = round(os.path.getsize(fallback_path) / 1024, 1)
                    self._try_open(fallback_path)
                    self._popup_save_success(fname, fallback_path, os.getcwd(), size_kb)
            except Exception as e2:
                self._popup("Save Failed", f"Could not save file.\n\nError:\n{str(e2)}", error=True)

    def _try_open(self, path):
        try:
            if sys.platform == "win32":
                os.startfile(path)
            elif sys.platform == "darwin":
                subprocess.Popen(["open", path])
            else:
                subprocess.Popen(["xdg-open", path])
        except Exception:
            pass

    def print_form(self, *_):
        d = self._data()

        def row(label, value):
            return f"<tr><td class='lbl'>{label}</td><td class='val'>{value or '&nbsp;'}</td></tr>"

        html = f"""<!DOCTYPE html>
<html><head><meta charset='UTF-8'><title>OPD Card — {d.get('name', 'Patient')}</title>
<style>
body {{ font-family: Arial, sans-serif; margin: 20px; color: #111; }}
h1 {{ text-align:center; color:#0d478d; margin-bottom:2px; font-size:18px; }}
h2 {{ text-align:center; color:#333; margin-top:2px; font-size:13px; }}
h3 {{ color:#b71c1c; border-bottom:1px solid #b71c1c; padding-bottom:3px; margin-top:14px; font-size:13px; }}
table {{ width:100%; border-collapse:collapse; margin-bottom:8px; }}
.lbl {{ font-weight:bold; width:40%; padding:5px 8px; background:#f0f4ff; border:1px solid #ccc; font-size:12px; }}
.val {{ padding:5px 8px; border:1px solid #ccc; font-size:12px; }}
.sig {{ position: fixed; bottom: 18px; right: 24px; font-weight: bold; font-size: 12px; color: #111; }}
@media print {{ body {{ margin: 10mm 12mm; }} .sig {{ position: fixed; bottom: 12mm; right: 12mm; font-size: 11pt; }} }}
</style></head><body>
<h1>JINDAL STEEL POWER LIMITED</h1>
<h2>Occupational Health Centre — Patratu (JH) &nbsp;|&nbsp; OPD CARD</h2>
<h3>Administrative Details</h3>
<table>{row("Regd. No.", d.get("regd_no"))}{row("Medical Officer", d.get("doctor"))}{row("Date &amp; Time", d.get("datetime"))}</table>
<h3>Patient Information</h3>
<table>{row("Name of Patient", d.get("name"))}{row("E Code / GP No.", d.get("ecode"))}{row("Age / Sex", d.get("age_sex"))}{row("Mobile No.", d.get("mobile"))}{row("Address / Deptt.", d.get("address"))}{row("Contractor", d.get("contractor"))}{row("Service", d.get("service"))}</table>
<h3>Vital Signs</h3>
<table>{row("Temperature", (d.get("temp") or "") + " °F")}{row("SPO2", (d.get("spo2") or "") + " %")}{row("Pulse Rate", (d.get("pr") or "") + " /min")}{row("Blood Pressure", (d.get("bp") or "") + " mmHg")}{row("Weight", (d.get("weight") or "") + " Kg")}</table>
<h3>Investigation / Notes</h3><p style='font-size:12px'>{d.get("investigation") or "—"}</p>
<h3>Remarks</h3><p style='font-size:12px'>{d.get("remarks") or "—"}</p>
<div class='sig'>Signature of Doctor: _________________________</div>
</body></html>"""

        html_path = os.path.join(os.getcwd(), "opd_preview.html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html)
        try:
            self._try_open(html_path)
            self._popup("Print Preview", "OPD Card opened in your browser.\n\nPress  Ctrl+P  to print\nor save as PDF.")
        except Exception as e:
            self._popup("Error", str(e), error=True)

    def _popup_save_success(self, fname, full_path, location_label, size_kb):
        content = BoxLayout(orientation="vertical", padding=dp(14), spacing=dp(8))
        content.add_widget(Label(text=f"[b]{fname}[/b]", markup=True, color=(0.05, 0.4, 0.1, 1),
                                  font_size=dp(13), halign="center", valign="middle",
                                  size_hint_y=None, height=dp(30), text_size=(dp(290), None)))
        content.add_widget(Label(text=f"[b]Location:[/b]  {location_label}", markup=True,
                                  color=(0.2, 0.2, 0.2, 1), font_size=dp(12), halign="left",
                                  valign="middle", size_hint_y=None, height=dp(24), text_size=(dp(290), None)))
        path_lbl = Label(text=full_path, color=(0.4, 0.4, 0.4, 1), font_size=dp(11),
                          halign="left", valign="top", size_hint_y=None, text_size=(dp(290), None))
        path_lbl.bind(texture_size=lambda w, v: setattr(w, "height", v[1] + dp(6)))
        content.add_widget(path_lbl)
        content.add_widget(Label(text=f"File saved successfully  •  {size_kb} KB",
                                  color=(0.05, 0.45, 0.15, 1), font_size=dp(12), halign="center",
                                  valign="middle", size_hint_y=None, height=dp(24), text_size=(dp(290), None)))
        content.add_widget(Label(text="(File opened automatically)", color=(0.5, 0.5, 0.5, 1),
                                  font_size=dp(11), halign="center", valign="middle",
                                  size_hint_y=None, height=dp(20), text_size=(dp(290), None)))
        btn = Button(text="OK", size_hint_y=None, height=dp(42), background_color=GREEN,
                     color=WHITE, font_size=dp(13), bold=True)
        content.add_widget(btn)
        pop = Popup(title="Word File Saved", content=content, size_hint=(0.88, 0.58),
                    title_color=GREEN, separator_color=GREEN)
        btn.bind(on_press=pop.dismiss)
        pop.open()

    def _popup(self, title, msg, error=False):
        color = RED if error else GREEN
        content = BoxLayout(orientation="vertical", padding=dp(14), spacing=dp(10))
        content.add_widget(Label(text=msg, color=(0.1, 0.1, 0.1, 1), font_size=dp(13),
                                  halign="center", valign="middle", text_size=(dp(280), None)))
        btn = Button(text="OK", size_hint_y=None, height=dp(42), background_color=color,
                     color=WHITE, font_size=dp(13), bold=True)
        content.add_widget(btn)
        pop = Popup(title=title, content=content, size_hint=(0.82, 0.42),
                    title_color=color, separator_color=color)
        btn.bind(on_press=pop.dismiss)
        pop.open()


# ══════════════════════════════════════════════════════════════════════════
#  SCREEN 2 : MEDICINE EXPIRY TRACKER
# ══════════════════════════════════════════════════════════════════════════
DB_NAME = "medicines.db"

def get_db_path():
    if getattr(sys, 'frozen', False):
        base_dir = os.path.dirname(os.path.abspath(sys.executable))
    else:
        base_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_dir, DB_NAME)

def init_db():
    conn = sqlite3.connect(get_db_path())
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS medicines (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            unit TEXT NOT NULL,
            qty INTEGER NOT NULL,
            expiry TEXT NOT NULL
        )
    """)
    conn.commit()
    conn.close()

def db_add_medicine(name, unit, qty, expiry_str):
    conn = sqlite3.connect(get_db_path())
    cur = conn.cursor()
    cur.execute("INSERT INTO medicines (name, unit, qty, expiry) VALUES (?, ?, ?, ?)",
                (name, unit, qty, expiry_str))
    conn.commit()
    conn.close()

def db_fetch_all():
    conn = sqlite3.connect(get_db_path())
    cur = conn.cursor()
    cur.execute("SELECT id, name, unit, qty, expiry FROM medicines ORDER BY expiry ASC")
    rows = cur.fetchall()
    conn.close()
    return rows

def db_delete_medicine(med_id):
    conn = sqlite3.connect(get_db_path())
    cur = conn.cursor()
    cur.execute("DELETE FROM medicines WHERE id = ?", (med_id,))
    conn.commit()
    conn.close()

LOW_STOCK_THRESHOLD = {"Strips": 10, "Tabs": 15}
EXPIRY_WARNING_DAYS = 15


class MedicineTrackerBody(BoxLayout):
    def __init__(self, **kwargs):
        super(MedicineTrackerBody, self).__init__(**kwargs)
        self.orientation = 'vertical'
        self.padding = 10
        self.spacing = 10

        init_db()
        self.medicines = []

        self.add_widget(Label(text="Medicine Expiry Tracker", font_size=24,
                               size_hint_y=None, height=40, color=BLACK))

        form = GridLayout(cols=2, spacing=10, size_hint_y=None, height=180)
        form.add_widget(Label(text="Medicine Name:", color=BLACK))
        self.name_input = TextInput(multiline=False)
        form.add_widget(self.name_input)

        form.add_widget(Label(text="Unit Type:", color=BLACK))
        self.unit_spinner = Spinner(text='Strips', values=('Strips', 'Tabs'))
        form.add_widget(self.unit_spinner)

        form.add_widget(Label(text="Quantity:", color=BLACK))
        self.qty_input = TextInput(multiline=False, input_filter='int')
        form.add_widget(self.qty_input)

        form.add_widget(Label(text="Expiry Date (DD-MM-YYYY):", color=BLACK))
        self.expiry_input = TextInput(multiline=False, hint_text="e.g., 31-12-2026")
        form.add_widget(self.expiry_input)

        self.add_widget(form)

        add_btn = Button(text="Add Medicine", size_hint_y=None, height=44,
                          background_color=(0.2, 0.6, 1, 1))
        add_btn.bind(on_press=self.add_medicine)
        self.add_widget(add_btn)

        self.alert_label = Label(text="Status: All clear", size_hint_y=None, height=60, color=(0, 0.6, 0, 1))
        self.alert_label.bind(size=self.alert_label.setter('text_size'))
        self.add_widget(self.alert_label)

        self.scroll_view = ScrollView()
        self.med_list_layout = GridLayout(cols=1, spacing=5, size_hint_y=None)
        self.med_list_layout.bind(minimum_height=self.med_list_layout.setter('height'))
        self.scroll_view.add_widget(self.med_list_layout)
        self.add_widget(self.scroll_view)

        self.load_from_db()

    def load_from_db(self):
        rows = db_fetch_all()
        self.medicines = []
        for row in rows:
            med_id, name, unit, qty, expiry_str = row
            expiry_date = datetime.strptime(expiry_str, "%d-%m-%Y").date()
            self.medicines.append({"id": med_id, "name": name, "unit": unit,
                                    "qty": int(qty), "expiry": expiry_date})
        self.refresh_list()

    def add_medicine(self, instance):
        name = self.name_input.text.strip()
        unit = self.unit_spinner.text
        qty = self.qty_input.text.strip()
        expiry_str = self.expiry_input.text.strip()

        if not name or not qty or not expiry_str:
            self.update_alert("Error: All fields are required!", is_error=True)
            return
        try:
            datetime.strptime(expiry_str, "%d-%m-%Y").date()
        except ValueError:
            self.update_alert("Error: Use DD-MM-YYYY format!", is_error=True)
            return

        db_add_medicine(name, unit, int(qty), expiry_str)
        self.name_input.text = ""
        self.qty_input.text = ""
        self.expiry_input.text = ""
        self.load_from_db()

    def refresh_list(self):
        self.med_list_layout.clear_widgets()
        today = datetime.now().date()
        expired_names, expiring_soon_names, low_stock_names = [], [], []

        for med in self.medicines:
            days_left = (med["expiry"] - today).days
            threshold = LOW_STOCK_THRESHOLD.get(med["unit"], 10)
            is_low_stock = med["qty"] < threshold

            if days_left <= 0:
                status_text = "[EXPIRED]"
                item_color = (1, 0.3, 0.3, 1)
                expired_names.append(med["name"])
            elif days_left <= EXPIRY_WARNING_DAYS:
                status_text = f"[Expires in {days_left} days]"
                item_color = (1, 0.5, 0, 1)
                expiring_soon_names.append(f"{med['name']} ({days_left}d)")
            else:
                status_text = f"[{days_left} days left]"
                item_color = BLACK

            stock_text = ""
            if is_low_stock:
                stock_text = " | LOW STOCK!"
                low_stock_names.append(f"{med['name']} ({med['qty']} {med['unit']})")
                item_color = (0.7, 0, 0.7, 1)

            row = BoxLayout(orientation='horizontal', size_hint_y=None, height=40, spacing=10)
            exp_display = med['expiry'].strftime("%d-%m-%Y")
            info_text = f"{med['name']} - {med['qty']} {med['unit']} | Exp: {exp_display} {status_text}{stock_text}"
            lbl = Label(text=info_text, color=item_color, halign='left', valign='middle', font_size=16)
            lbl.bind(size=lbl.setter('text_size'))
            row.add_widget(lbl)

            del_btn = Button(text="Delete", size_hint_x=None, width=80, background_color=(0.8, 0.2, 0.2, 1))
            del_btn.bind(on_press=lambda inst, med_id=med["id"]: self.delete_medicine(med_id))
            row.add_widget(del_btn)
            self.med_list_layout.add_widget(row)

        self.build_alert_message(expired_names, expiring_soon_names, low_stock_names)

    def build_alert_message(self, expired_names, expiring_soon_names, low_stock_names):
        messages = []
        if expired_names:
            messages.append(f"EXPIRED: {', '.join(expired_names)}")
        if expiring_soon_names:
            messages.append(f"Expiring within {EXPIRY_WARNING_DAYS} days: {', '.join(expiring_soon_names)}")
        if low_stock_names:
            messages.append(f"Low Stock: {', '.join(low_stock_names)}")

        if messages:
            self.update_alert(" || ".join(messages), is_error=True)
        else:
            self.update_alert("Status: System OK. No expired/low-stock medicines.", is_error=False)

    def delete_medicine(self, med_id):
        db_delete_medicine(med_id)
        self.load_from_db()

    def update_alert(self, text, is_error=False):
        self.alert_label.text = text
        self.alert_label.color = (0.8, 0, 0, 1) if is_error else (0, 0.6, 0, 1)

    def on_start_check(self):
        today = datetime.now().date()
        urgent = []
        for med in self.medicines:
            days_left = (med["expiry"] - today).days
            threshold = LOW_STOCK_THRESHOLD.get(med["unit"], 10)
            if days_left <= 0:
                urgent.append(f"{med['name']} - EXPIRED")
            elif days_left <= EXPIRY_WARNING_DAYS:
                urgent.append(f"{med['name']} - expires in {days_left} days")
            if med["qty"] < threshold:
                urgent.append(f"{med['name']} - LOW STOCK ({med['qty']} {med['unit']})")

        if urgent:
            content = BoxLayout(orientation="vertical", padding=dp(14), spacing=dp(10))
            with content.canvas.before:
                Color(0.85, 0.93, 1, 1)
                content._bg_rect = Rectangle(pos=content.pos, size=content.size)
            content.bind(pos=lambda w, v: setattr(w._bg_rect, "pos", v),
                         size=lambda w, v: setattr(w._bg_rect, "size", v))

            msg_lbl = Label(text="\n".join(urgent), color=BLACK, font_size=dp(13),
                             halign="left", valign="top", text_size=(dp(280), None))
            content.add_widget(msg_lbl)
            close_btn = Button(text="OK", size_hint_y=None, height=dp(42),
                                background_color=(0.2, 0.6, 1, 1), color=WHITE,
                                font_size=dp(13), bold=True)
            content.add_widget(close_btn)

            popup = Popup(title="Medicine Alerts", content=content, size_hint=(0.85, 0.6),
                           title_color=BLACK, separator_color=(0.2, 0.6, 1, 1))
            close_btn.bind(on_press=popup.dismiss)
            popup.open()


class MedicineScreen(Screen):
    def __init__(self, **kw):
        super().__init__(**kw)
        root = BoxLayout(orientation="vertical")
        root.add_widget(build_nav_bar(self))
        self.body = MedicineTrackerBody()
        root.add_widget(self.body)
        self.add_widget(root)


# ══════════════════════════════════════════════════════════════════════════
#  SCREEN 3 : MEDBOT CHATBOT
# ══════════════════════════════════════════════════════════════════════════
MEDICINE_DB = {
    "fever": ["Paracetamol 650mg", "Ibuprofen 400mg", "Dolo 650mg"],
    "headache": ["Saridon", "Paracetamol 500mg", "Aspirin"],
    "pain": ["Tramadol", "Diclofenac Gel", "Zerodol-P"],
    "allergy": ["Cetirizine 10mg", "Allegra 120mg", "Levocetirizine"],
    "acidity": ["Pantocid 40mg", "Pan-D", "Omeprazole", "Digene Syrup"],
    "cough": ["Ascoril LS", "Benadryl Cough Syrup", "Ambroxol"],
    "cold": ["Solvin Cold", "Sinarest", "Cheston Cold"],
    "diarrhea": ["ORS Powder", "Loperamide", "Metrogyl 400mg"],
    "weakness": ["Zincovit", "Neurobion Forte", "Limcee Vitamin C"],
}

ALIASES = {"bp": "high blood pressure", "sugar": "blood sugar",
           "loose motion": "diarrhea", "gas": "acidity"}


class ChatBubble(BoxLayout):
    def __init__(self, text, is_bot=True, **kwargs):
        super().__init__(**kwargs)
        self.orientation = "horizontal"
        self.size_hint_y = None
        self.padding = [dp(10), dp(5)]
        self.spacing = dp(8)

        bubble_color = get_color_from_hex("#E3F2FD") if is_bot else get_color_from_hex("#1E3A8A")
        text_color = get_color_from_hex("#0D1B2A") if is_bot else get_color_from_hex("#FFFFFF")

        avatar = Label(text="Bot" if is_bot else "You", font_size=dp(12), size_hint=(None, None), size=(dp(34), dp(28)))

        self.msg = Label(text=text, markup=True, color=text_color, font_size='14sp',
                          text_size=(dp(260), None), halign="left", valign="top")
        self.msg.bind(texture_size=self.msg.setter("size"))

        bubble_wrap = BoxLayout(size_hint=(None, None), padding=[dp(12), dp(10)])
        with bubble_wrap.canvas.before:
            Color(*bubble_color)
            bubble_wrap.bg_rect = RoundedRectangle(pos=bubble_wrap.pos, size=bubble_wrap.size, radius=[12])
        bubble_wrap.bind(
            pos=lambda *a: setattr(bubble_wrap.bg_rect, "pos", bubble_wrap.pos),
            size=lambda *a: setattr(bubble_wrap.bg_rect, "size", bubble_wrap.size),
        )
        bubble_wrap.add_widget(self.msg)
        self.msg.bind(size=lambda *a: setattr(bubble_wrap, "size", (self.msg.width + dp(24), self.msg.height + dp(20))))

        if is_bot:
            self.add_widget(avatar)
            self.add_widget(bubble_wrap)
            self.add_widget(Widget())
        else:
            self.add_widget(Widget())
            self.add_widget(bubble_wrap)
            self.add_widget(avatar)

        self.bind(minimum_height=self.setter("height"))
        Clock.schedule_once(lambda dt: self._fix_height(), 0.05)

    def _fix_height(self):
        heights = [c.height for c in self.children if c.height > 0]
        self.height = (max(heights) if heights else dp(40)) + dp(12)


class MedBotBody(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.orientation = 'vertical'
        self.spacing = 10
        self.padding = 15
        self._processing = False

        self.add_widget(Label(text="[b]MedBot AI Assistant[/b]", markup=True, font_size='18sp',
                               color=get_color_from_hex('#1E3A8A'), size_hint_y=None, height=40))

        self.scroll = ScrollView(size_hint=(1, 1), do_scroll_x=False)
        self.chat_box = BoxLayout(orientation="vertical", size_hint_y=None, spacing=10, padding=[5, 5])
        self.chat_box.bind(minimum_height=self.chat_box.setter('height'))
        self.scroll.add_widget(self.chat_box)
        self.add_widget(self.scroll)

        quick_grid = GridLayout(cols=3, size_hint_y=None, height=100, spacing=6)
        symptoms_list = ["Fever", "Headache", "Cough", "Cold", "Acidity", "Diarrhea"]
        for s in symptoms_list:
            btn = Button(text=s, font_size='13sp', bold=True, background_normal='',
                         background_color=get_color_from_hex('#E3F2FD'), color=get_color_from_hex('#1E3A8A'))
            btn.bind(on_release=lambda b, sym=s: self._process_query(sym))
            quick_grid.add_widget(btn)
        self.add_widget(quick_grid)

        input_row = BoxLayout(orientation='horizontal', size_hint_y=None, height=50, spacing=6)
        self.text_input = TextInput(hint_text="Type symptoms here...", multiline=False,
                                     font_size='14sp', padding=[12, 14, 12, 12])
        self.text_input.bind(on_text_validate=self._on_text_submit)
        btn_send = Button(text="Send", size_hint_x=None, width=75, bold=True,
                           background_normal='', background_color=get_color_from_hex('#1E3A8A'))
        btn_send.bind(on_release=self._on_text_submit)
        input_row.add_widget(self.text_input)
        input_row.add_widget(btn_send)
        self.add_widget(input_row)

        Clock.schedule_once(lambda dt: self._bot_say(
            "Hello! Main aapka healthcare assistant hoon. Mujhe apne symptoms bataiye "
            "(jaise Fever, Headache, ya Cough) taaki main medicine suggest kar sakoon."
        ), 0.2)

    def _on_text_submit(self, *args):
        query = self.text_input.text.strip()
        if query:
            self.text_input.text = ""
            self._process_query(query)

    def _process_query(self, query):
        if self._processing:
            return
        self._processing = True
        self.chat_box.add_widget(ChatBubble(text=query, is_bot=False))

        text = query.lower().strip()
        for alias, canonical in ALIASES.items():
            if alias in text:
                text = text.replace(alias, canonical)

        found = False
        response = "[b][color=#10B981]Suggested Medications:[/color][/b]\n"
        for symptom, meds in MEDICINE_DB.items():
            if symptom in text:
                found = True
                for med in meds:
                    response += f"\n- [b]{med}[/b]"

        if found:
            self.chat_box.add_widget(ChatBubble(text=response, is_bot=True))
        else:
            self.chat_box.add_widget(ChatBubble(
                text=f"Symptom '[b]{query}[/b]' ke liye database mein koi records nahi mile. "
                     f"Kripya dusra symptom check karein.", is_bot=True))

        Clock.schedule_once(lambda dt: setattr(self.scroll, 'scroll_y', 0), 0.1)
        self._processing = False

    def _bot_say(self, text):
        self.chat_box.add_widget(ChatBubble(text=text, is_bot=True))


class MedBotScreen(Screen):
    def __init__(self, **kw):
        super().__init__(**kw)
        root = BoxLayout(orientation="vertical")
        root.add_widget(build_nav_bar(self))
        self.body = MedBotBody()
        root.add_widget(self.body)
        self.add_widget(root)


# ══════════════════════════════════════════════════════════════════════════
#  APP ENTRY
# ══════════════════════════════════════════════════════════════════════════
class JindalOHCApp(App):
    def build(self):
        self.title = "Jindal OHC — Unified App"
        sm = ScreenManager()
        sm.add_widget(OPDScreen(name="opd"))
        sm.add_widget(MedicineScreen(name="medicine"))
        sm.add_widget(MedBotScreen(name="medbot"))
        sm.current = "opd"
        self.sm = sm
        return sm

    def on_start(self):
        # Medicine screen ke urgent alerts ka popup dikhayein
        self.sm.get_screen("medicine").body.on_start_check()


if __name__ == "__main__":
    JindalOHCApp().run()
